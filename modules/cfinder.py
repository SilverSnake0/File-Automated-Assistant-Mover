import PyPDF2
import pandas as pd
import docx
import os

def preview(file_path):
    # Checks if the file has a txt, word, or pdf extension then asks if they would like to preview the file
    if file_path.endswith('.txt'):
        ask_question = input('\n\nThis file extension may be supported for preview.\nWould you like attempt to preview this file? ("yes" for yes or "no" for no): \n\n')
        if ask_question.lower() == 'yes':
            # Reads the contents of the text file
            with open(file_path, 'r') as f:
                contents = f.read()
            
            # Splits the contents into chunks
            chunk_size = 500
            chunks = [contents[i:i+chunk_size] for i in range(0, len(contents), chunk_size)]
            
            # Previews the contents, chunk by chunk
            for chunk in chunks:
                print(chunk)
                response = input('\n\nWould you like to continue previewing the file? (Press enter to continue or type "no" for no): \n\n')
                if response.lower() == 'no':
                    break
    
    elif file_path.endswith('.docx'):
        ask_question = input('\n\nThis file extension may be supported for preview.\nWould you like attempt to preview this file? ("yes" for yes or "no" for no): \n\n')
        if ask_question.lower() == 'yes':
            # Opens the docx file
            doc = docx.Document(file_path)
            
            # Previews the contents, paragraph by paragraph
            for i, paragraph in enumerate(doc.paragraphs):
                print(f"Paragraph {i+1}: {paragraph.text}")
                response = input('\n\nWould you like to continue previewing the file? (Press enter to continue or type "no" for no): \n\n')
                if response.lower() == 'no':
                    break
    
    elif file_path.endswith('.pdf'):
        ask_question = input('\n\nThis file extension may be supported for preview.\nWould you like attempt to preview this file? ("yes" for yes or "no" for no): \n\n')
        if ask_question.lower() == 'yes':
            # Open the pdf file
            with open(file_path, 'rb') as f:
                pdf = PyPDF2.PdfFileReader(f)
            
            # Previews the contents, page by page
            for i in range(pdf.getNumPages()):
                page = pdf.getPage(i)
                print(page.extractText())
                response = input('\n\nWould you like to continue previewing the file? (Press enter to continue or type "no" for no): \n\n')
                if response.lower() == 'no':
                    break

def search_file_content(file_path, string):
    # Checks the file extension
    if file_path.endswith('.txt'):
        # Opens the file and searches for the string
        with open(file_path, 'r', encoding='utf-8') as f:
            content = f.read()
            return string in content
    elif file_path.endswith('.pdf'):
        # Uses PyPDF2 to search for the string in the PDF
        with open(file_path, 'rb') as f:
            pdf = PyPDF2.PdfReader(f)
            # Iterates through all pages in the PDF
            for page in pdf.pages:
                # Extracts the text from the page
                text = page.extractText()
                # Searches for the string in the text
                if string in text:
                    return True
            return False
    elif file_path.endswith('.docx'):
        # Uses docx to search for the string in the Word file
        doc = docx.Document(file_path)
        # Iterates through all paragraphs in the document
        for paragraph in doc.paragraphs:
            # Extracts the text from the paragraph
            text = paragraph.text
            # Searches for the string in the text
            if string in text:
                return True
        return False
    elif file_path.endswith('.xlsx'):
        # Reads the Excel file into a pandas DataFrame
        df = pd.read_excel(file_path)
        # Finds rows where any cell value contains the string
        df_filtered = df[df.apply(lambda x: string in x, axis=1)]
        # Returns True if there are any matching rows
        return not df_filtered.empty
    else:
        return False

def search_by_content(files, string, source_dir, target, matching_files): 
    if target == 'Matching Files':
        files = matching_files
    print(f'Current Target: {target}\n')
    # Searches through each file for the string
    for file in files:
        try:
            print(f'Searching through file content: {file}')
            file_path = os.path.join(source_dir, file)
            if search_file_content(file_path, string):
                if file_path not in matching_files:
                    print(f'Found {string} in {file}!')
                    matching_files.append(file_path)
        except Exception as e:
            # Prints the error message if an error occurs
            print(f'An error occurred while reading {file} file: {e}')
            continue
    print(f'Search has finished.')
    file_names = [os.path.basename(path) for path in matching_files]
    print(f'Matching Files: {file_names}')
    return matching_files